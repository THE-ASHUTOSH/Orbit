"""Build the Orbit guide as a Word document.

    pip install python-docx
    python3 scripts/make-guide-docx.py docs/Orbit-Guide.docx

The whole guide, not just the usage half: setting it up, using it, running it,
and fixing it. Same ground as docs/team-guide.md, for people who would rather be
handed a document than read a repository - change both, or neither.

Every UI label quoted here was read out of the app's source, and every command
was run before it was written down.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

OUT = sys.argv[1]
doc = Document()

# --- document setup ---------------------------------------------------------
p = doc.core_properties
p.title = 'Orbit - Complete Guide'
p.subject = 'Setting up, using and running the shared team browser'
p.author = 'the-ashutosh'
p.comments = 'Setup, usage, administration and troubleshooting for Orbit.'

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8)
    s.left_margin = s.right_margin = Inches(0.9)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def callout(label, text, fill='FFF3CD'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(10)
    para.paragraph_format.left_indent = Inches(0.12)
    r = para.add_run(f'{label}  ')
    r.bold = True
    para.add_run(text)
    shade(para, fill)
    return para


def bullets(items, style='List Bullet'):
    for it in items:
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(3)
        # **bold** segments
        for i, chunk in enumerate(it.split('**')):
            run = para.add_run(chunk)
            run.bold = i % 2 == 1


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    for style in ('Light List Accent 1', 'Light Grid Accent 1', 'Table Grid'):
        try:
            t.style = style
            break
        except KeyError:
            continue
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            for j, chunk in enumerate(str(val).split('**')):
                run = para.add_run(chunk)
                run.bold = j % 2 == 1
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_numbers():
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Orbit - Complete Guide    ').font.size = Pt(8)
    run = para.add_run()
    run.font.size = Pt(8)
    for kind, text in (('begin', None), (None, 'PAGE'), ('end', None)):
        if text:
            el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve'); el.text = text
        else:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), kind)
        run._r.append(el)




def part(number, title, blurb):
    """A major division, each starting on its own page."""
    if part.started:
        doc.add_page_break()
    part.started = True
    head = doc.add_heading(f'Part {number} - {title}', level=1)
    head.paragraph_format.space_before = Pt(0)
    para = doc.add_paragraph()
    r = para.add_run(blurb)
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    para.paragraph_format.space_after = Pt(12)


part.started = False


def steps(items):
    for text in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(5)
        for i, chunk in enumerate(text.split('**')):
            run = para.add_run(chunk)
            run.bold = i % 2 == 1


def code(lines):
    """A command block: monospace, shaded, tight."""
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = Inches(0.2)
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        shade(para, 'F2F2F2')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# --- title ------------------------------------------------------------------
doc.add_heading('Orbit', level=0)
sub = doc.add_paragraph()
r = sub.add_run('Set it up, use it, run it')
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.paragraph_format.space_after = Pt(14)

doc.add_paragraph(
    'Orbit is one real browser running on one computer, shared with everyone on your network. '
    'Each person opens it in their own browser and works in their own tabs - several people '
    'browsing at the same time, independently, inside the same shared browser.'
)
doc.add_paragraph(
    'It is not screen sharing. Nobody has to present, and you are not watching someone else move '
    'their mouse. The thing that IS shared is logins: sign in to a website once and everyone '
    'using Orbit is signed in to it.'
)

doc.add_heading('Contents', level=2)
CONTENTS = [
    ('Part 1 - Set it up', 'for whoever hosts it', [
        '1.1  What you need',
        '1.2  Install Docker',
        '1.3  Get Orbit and start it',
        '1.4  Find the address everyone will use',
        '1.5  Add your team',
    ]),
    ('Part 2 - Using Orbit', 'for everyone - hand this part to your team', [
        '2.1  Getting in',
        '2.2  The window',
        '2.3  Tabs belong to whoever opens them',
        '2.4  Keyboard shortcuts',
        '2.5  Capture keyboard',
        '2.6  Tabs: everything you can do',
        '2.7  Zoom',
        '2.8  Full screen',
        '2.9  Bookmarks and history',
        '2.10  Downloads: getting a file onto your own computer',
        '2.11  Uploading a file to a page',
        '2.12  Right-click menu',
        '2.13  Extensions',
        '2.14  Who else is here, and how it is doing',
        '2.15  The one thing to remember',
    ]),
    ('Part 3 - Running it', 'settings, extensions, outside access, upkeep', [
        '3.1  Settings, in plain words',
        '3.2  Installing extensions',
        '3.3  Access from outside the network, with ngrok',
        '3.4  Everyday commands',
        '3.5  Where the data lives, and backups',
    ]),
    ('Part 4 - When something goes wrong', 'symptoms and fixes, for both', [
        '4.1  Quick table',
        '4.2  Chrome on a Mac says "This site can\'t be reached"',
        '4.3  Other devices cannot reach it',
        '4.4  ngrok: the link opens but I cannot sign in',
        '4.5  Who to ask',
    ]),
]
for heading, blurb, items in CONTENTS:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    para.add_run(heading).bold = True
    r = para.add_run(f'  -  {blurb}')
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    for item in items:
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Inches(0.3)
        line.paragraph_format.space_after = Pt(0)
        line.add_run(item).font.size = Pt(9.5)

callout('New to this?',
        'If somebody has already set Orbit up and just sent you a link, you only need Part 2.',
        fill='E7F3EB')

# =============================================================================
part(1, 'Set it up', 'About five minutes, most of it downloading. Do this once, on one computer.')
# =============================================================================

doc.add_heading('1.1  What you need', level=2)
doc.add_paragraph(
    'One computer runs Orbit and does all the work. Everyone else needs nothing but a browser.'
)
table(['', 'The host computer'], [
    ['Operating system', 'macOS, Windows or Linux'],
    ['Software', 'Docker (step 1.2), plus git and openssl - already on macOS and Linux'],
    ['Free capacity', '4 processor cores and 4 GB of memory is fine; 6 and 6 is comfortable'],
    ['Network', 'on the same Wi-Fi or wired network as your team'],
    ['Disk', 'about 2 GB for the browser, plus room for its saved data'],
], widths=[1.6, 4.9])
callout('Keep it awake',
        'Leave the host computer on while people are using it, and plug a laptop in. Sending that '
        'machine to sleep stops the browser for everybody.')

doc.add_heading('1.2  Install Docker', level=2)
doc.add_paragraph('Docker is the tool that runs Orbit in a self-contained box.')
bullets([
    '**macOS or Windows:** download Docker Desktop from docker.com and install it. Open it once, '
    'so it is running - you will see a small whale icon in the menu bar or system tray.',
    '**Linux:** install Docker Engine using your distribution\'s instructions.',
])
doc.add_paragraph('Check it worked. Open Terminal (macOS) or PowerShell (Windows) and run:')
code(['docker --version'])
doc.add_paragraph(
    'You should see a version number. "command not found" means Docker is not installed, or is '
    'not running yet.'
)

doc.add_heading('1.3  Get Orbit and start it', level=2)
code([
    'git clone https://github.com/THE-ASHUTOSH/Orbit.git',
    'cd Orbit',
    './orbit up',
])
doc.add_paragraph(
    'The first run builds the browser on your machine, so give it a few minutes; later starts take '
    'seconds. That one command does all of the setup for you:'
)
steps([
    'Starts Docker if it is not already running.',
    'Notices there is no **.env** yet and creates one - so there is no file to write by hand.',
    'Generates a random **SESSION_SECRET** for you. You never see it or choose it.',
    'Asks you to **choose an admin password**, twice, to catch typos. Eight characters or more.',
    'Fills in the rest of the settings with the values this project actually runs on.',
    'Fetches the default extension and **asks whether you want it installed** - name, version and '
    'what it can access. Answer once; it is remembered. Section 3.2 has the detail.',
    'Builds the browser, waits until it is genuinely healthy, and prints the two addresses - one '
    'for this machine, one for everyone else.',
])
doc.add_paragraph('It looks like this:')
code([
    '> no .env yet - creating one',
    '  choose an admin password (8+ chars):',
    '  again:',
    '> wrote .env (chmod 600) - the same settings this project runs on',
    '> building and starting',
    '> waiting for the browser to come up',
    '  open on this machine:  http://127.0.0.1:3030',
    '  open on the LAN:       http://192.168.1.42:3030',
])
doc.add_paragraph(
    'The password you type there is for the user **admin**. That is the only value you choose, and '
    'the only one you need to remember. Everything else is optional and can be changed later - see '
    'Part 3.'
)
callout('Why the admin password matters',
        'Orbit shares one set of website logins across everyone who uses it, so whoever gets in can '
        'see every account that browser is signed into. Treat it like the key to a shared office, '
        'not a test login. You can change it later, and add or remove people, in the Admin panel.',
        fill='FBE2E2')
para = doc.add_paragraph()
para.add_run('You also get the helper commands, which the rest of this guide uses:').bold = True
table(['Command', 'Does'], [
    ['./orbit up', 'Start it, or apply settings you changed'],
    ['./orbit status', 'Is it healthy? Plus live figures and the addresses'],
    ['./orbit url', 'Just the addresses, to paste to your team'],
    ['./orbit logs', 'Follow the log when something looks wrong'],
    ['./orbit user <name> [role]', 'Create an account from the terminal; prompts for the password'],
    ['./orbit users', 'List everyone'],
    ['./orbit restart', 'Restart the browser, keeping everyone signed in'],
    ['./orbit down', 'Stop it, keep all data'],
    ['./orbit backup [file]', 'Save everything to one archive'],
    ['./orbit env', 'Show the settings in use, with secrets hidden'],
], widths=[2.25, 4.25])
doc.add_paragraph(
    'You are set up. Section 1.4 finds the address to share; 1.5 adds your team.'
)

doc.add_heading('1.4  Find the address everyone will use', level=2)
doc.add_paragraph(
    'You do not have to work this out. ./orbit up prints both addresses when it finishes - the one '
    'for this machine, and the one everybody else on the network uses:'
)
code([
    '  open on this machine:  http://127.0.0.1:3030',
    '  open on the LAN:       http://192.168.1.42:3030',
])
doc.add_paragraph(
    'The second one is what you send your team. To see them again at any time, without restarting '
    'anything:'
)
code(['./orbit url', './orbit status      # the same addresses, plus health and live figures'])
doc.add_paragraph(
    'If you ever need to find that address by hand - the terminal window is long gone, or you are '
    'checking it changed - each system reports it differently:'
)
table(['Host computer', 'Command', 'Looks like'], [
    ['macOS', 'ipconfig getifaddr en0', '192.168.1.42'],
    ['Windows', 'ipconfig  (look for "IPv4 Address")', '192.168.1.42'],
    ['Linux', 'hostname -I', '192.168.1.42'],
], widths=[1.2, 3.4, 1.9])
doc.add_paragraph('Whichever way you got it, everyone opens that address with :3030 on the end:')
code(['http://192.168.1.42:3030'])
callout('It can change',
        'That address belongs to the network, not to Orbit: it can change when the host computer '
        'reconnects to Wi-Fi or is restarted. If the link stops working for your team, run '
        './orbit url again and send the new one.')
doc.add_paragraph('Sign in as admin, with the password you chose in step 1.4.')
callout('Mac users on Chrome',
        'If that address gives "This site can\'t be reached - ERR_ADDRESS_UNREACHABLE" while it '
        'works in Safari, it is a macOS permission, not a network fault. Section 4.2 has the fix.')

doc.add_heading('1.5  Add your team', level=2)
doc.add_paragraph('Accounts are created in the app, not on the command line.')
steps([
    'Sign in as **admin**.',
    'Click the **⋮ Menu** button at the top right.',
    'Choose **Admin panel**.',
    'Open the **Users** section and add each person: a username, a password and a role.',
])
doc.add_paragraph('Then send each person the address, their username and their password.')
doc.add_paragraph(
    'The terminal works too: ./orbit user priya makes a normal user and '
    'asks for their password, ./orbit user sam viewer makes a watcher, and ./orbit users lists '
    'everyone. Same accounts either way.'
)
table(['Role', 'Can do'], [
    ['admin', 'Everything: manage users, install extensions, restart the browser, control any tab.'],
    ['user', 'Open tabs, browse, type. The normal role for a teammate.'],
    ['viewer', 'Watch only. Can never type or click inside a page.'],
], widths=[1.2, 5.3])
doc.add_paragraph(
    'The other Admin panel sections are Tabs (see and close any tab), Extensions (install for '
    'everyone), Browser (restart it) and Shared browser state (what it is signed into).'
)

# =============================================================================
part(2, 'Using Orbit', 'For everyone. Every feature, and the three that behave differently from a normal browser: tab ownership, downloads and extensions.')
# =============================================================================

doc.add_heading('2.1  Getting in', level=2)
doc.add_paragraph('Ask whoever set Orbit up for three things:')
bullets([
    '**The address** - something like http://192.168.1.42:3030, or an https:// link if you are '
    'outside the office.',
    '**Your username.**',
    '**Your password.**',
])
doc.add_paragraph(
    'Open the address in any browser, on any device - laptop, phone or tablet. Enter your username '
    'and password and click Login. Nothing to install.'
)
callout('Note',
        'You must be on the same Wi-Fi or network as the host computer, unless you were given an '
        'https:// link. Mobile data cannot reach a 192.168.x.x address.')

doc.add_heading('2.2  The window', level=2)
table(['Where', 'What it is'], [
    ['Top strip', 'The tabs. Click one to switch, + for a new one, x to close.'],
    ['Toolbar below it', 'Back, Forward, Reload, the address bar, and a star to bookmark the page.'],
    ['⋮ button, top right', 'The menu: zoom, full screen, bookmarks, history, downloads, extensions, appearance, sign out.'],
    ['Bottom strip', 'Who else is online, and how your connection is doing.'],
], widths=[1.6, 4.9])
doc.add_paragraph(
    'The address bar behaves as you expect: type a web address, or type words to search. It also '
    'suggests pages visited before.'
)

doc.add_heading('2.3  Tabs belong to whoever opens them', level=2)
doc.add_paragraph(
    'This is the one rule worth understanding. When you open a tab, that tab is yours. Others can '
    'see it but cannot type or click in it - so nobody fills in your form or clicks Send on your '
    'email by accident.'
)
doc.add_paragraph('To use a tab somebody else opened:')
steps([
    'Open the tab. The toolbar shows a button: **"Ask <name> for control"**.',
    'Click it. It changes to **"Asked <name>..."** while you wait.',
    'The owner sees a box in the corner of that same tab - **"Can I drive this tab?"** - with '
    '**Give control** and **Keep it to myself**.',
    'If they give control, you can type and click in that tab too.',
])
doc.add_paragraph(
    'Nothing is lost if they say no or are away from their desk: you can still read the page, and '
    'you can always open your own tab to the same site.'
)
callout('Tip',
        'Working on something you do not want touched? Keep it in a tab you opened. That is '
        'already private to you, with no extra steps.', fill='E7F3EB')

doc.add_heading('2.4  Keyboard shortcuts', level=2)
doc.add_paragraph(
    'Orbit\'s own shortcuts use Option on a Mac and Alt on Windows and Linux. They deliberately '
    'avoid Ctrl and Command so they never clash with your own browser\'s shortcuts.'
)
table(['Press', 'Does'], [
    ['Option / Alt + T', 'New tab'],
    ['Option / Alt + W', 'Close the tab'],
    ['Option / Alt + D', 'Jump to the address bar'],
    ['Option / Alt + F', 'Full screen'],
    ['Option / Alt + K', 'Capture keyboard (section 2.5)'],
    ['Option / Alt + left / right arrow', 'Back / Forward'],
    ['Option / Alt + 1 ... 9', 'Switch to that tab'],
], widths=[2.3, 4.2])
para = doc.add_paragraph()
para.add_run('Copy and paste are normal. ').bold = True
para.add_run(
    'Command+C / Command+V on a Mac, Ctrl+C / Ctrl+V on Windows. They work both ways: copy in '
    'Orbit and paste into your own documents, or copy on your computer and paste into a page '
    'inside Orbit.'
)

doc.add_heading('2.5  Capture keyboard', level=2)
doc.add_paragraph(
    'Some keys belong to your own browser. Press Ctrl+T (or Command+T) and YOUR browser opens a '
    'tab - Orbit never sees it. Turn on Capture keyboard and those keys go to Orbit instead.'
)
bullets([
    'Turn it on from the ⋮ menu, or with **Option / Alt + K**.',
    'It applies to **that one tab only**, and only while you are looking at it.',
    'A small badge shows when it is on. Turn it off the same way.',
])
doc.add_paragraph(
    'Use it when you want a page to feel like a normal browser tab. Leave it off the rest of the '
    'time so your own shortcuts keep working.'
)

doc.add_heading('2.6  Tabs: everything you can do', level=2)
table(['Action', 'How'], [
    ['New tab', 'The + on the tab strip, ⋮ menu, or Option / Alt + T'],
    ['Close a tab', 'The x on the tab, or Option / Alt + W'],
    ['Reopen the tab you just closed', 'Option / Alt + Shift + T'],
    ['Switch tabs', 'Click one, or Option / Alt + 1 to 9 for the first nine'],
    ['Duplicate a tab', '⋮ menu, Duplicate tab - opens the same page in a second tab, which is yours'],
    ['Rename a tab', 'Double-click the tab and type a name. Useful when four tabs all say "Dashboard"'],
    ['See whose tab it is', 'Hover it - the tooltip shows the page title, the address, and who opened it'],
], widths=[2.4, 4.1])
doc.add_paragraph(
    'A tab loading shows a small spinning ring in place of its icon, and a thin progress line over '
    'the page, so you can tell a slow site from a stuck one.'
)

doc.add_heading('2.7  Zoom', level=2)
doc.add_paragraph(
    'Open the ⋮ menu - the zoom control is at the top. Set a percentage, or click the percentage '
    'itself to snap back to 100%.'
)
bullets([
    'Zooming **out** shows more of the page, at smaller text - useful for wide dashboards.',
    'Zooming **in** enlarges everything, like a normal browser.',
    'The control also shows the **streamed resolution** at that zoom, which is what the page is '
    'actually being rendered at.',
])
doc.add_paragraph(
    'Zoom is per tab, and it is shared: someone else looking at that tab sees the same zoom, '
    'because it changes how the page is rendered rather than how your screen displays it.'
)

doc.add_heading('2.8  Full screen', level=2)
doc.add_paragraph(
    'Option / Alt + F, or ⋮ menu, Full screen. The page fills your whole screen and Orbit\'s tab '
    'strip and toolbar stay visible, so you can still switch tabs and use the address bar. Press '
    'the same keys again, or Escape, to come back.'
)

doc.add_heading('2.9  Bookmarks and history', level=2)
bullets([
    '**Bookmark the page you are on:** the star in the toolbar. Click it again to remove it.',
    '**See them all:** ⋮ menu, Bookmarks. Each entry opens in a new tab; the x removes it.',
    '**History:** ⋮ menu, History - with a search box, so you can find a page somebody visited '
    'last week by typing part of its name.',
])
callout('Both are shared',
        'Bookmarks and history belong to the shared browser, not to you: everyone sees the same '
        'lists. Bookmark the things the team needs, and remember that what you visit is visible to '
        'your colleagues.')
doc.add_paragraph(
    'The address bar also suggests from that history as you type, so a page anyone has visited is '
    'usually two or three keystrokes away.'
)

doc.add_heading('2.10  Downloads: getting a file onto your own computer', level=2)
doc.add_paragraph(
    'This is the part that surprises people, so it is worth understanding. When you download '
    'something, the file is downloaded by the shared browser - onto the host computer, not yours. '
    'Getting it to your own machine is a second, deliberate step.'
)
steps([
    'Click a download link in a page as normal. A notice appears when it finishes.',
    'Open **⋮ menu, Downloads**. The file is listed there, with its size.',
    'Click **Save** next to it. That is the point at which the file lands in your own Downloads '
    'folder, through your own browser.',
    'When nobody needs it any more, click the **delete** button next to it to remove it from the '
    'host computer.',
])
doc.add_paragraph(
    'The Refresh button re-reads the list, which is handy if a download finished while the panel '
    'was open.'
)
callout('Downloads are shared too',
        'Anything downloaded sits on the host computer until somebody deletes it, and everyone can '
        'see it and save a copy. Delete anything confidential once you have saved it, and do not '
        'download anything that should not be shared with the team.')

doc.add_heading('2.11  Uploading a file to a page', level=2)
doc.add_paragraph(
    'The reverse also works. When a page asks for a file - an attach button, a profile picture, a '
    'document upload - Orbit shows a file picker on **your** device, and you choose a file from '
    'your own computer.'
)
doc.add_paragraph(
    'Behind the scenes the file is sent to the host computer first, and then handed to the page, '
    'because the shared browser cannot reach into your filesystem. You do not have to do anything '
    'about that, but it explains the brief pause on a large file, and it means the file is briefly '
    'on the shared machine. If the page accepts several files, you can pick several.'
)

doc.add_heading('2.12  Right-click menu', level=2)
doc.add_paragraph(
    'Right-clicking inside a page gives you Orbit\'s own menu, not your browser\'s. What it '
    'offers depends on what you clicked:'
)
table(['On a link', 'On an image'], [
    ['Open link in new tab', 'Open image in new tab'],
    ['Copy link address', 'Copy image address'],
], widths=[3.2, 3.3])
doc.add_paragraph(
    'Copied addresses go to your own clipboard, ready to paste anywhere. For copying text, just '
    'select it and use Command+C / Ctrl+C as usual.'
)
doc.add_paragraph(
    'A tab you opened from a link this way belongs to you, like any other tab you open.'
)

doc.add_heading('2.13  Extensions', level=2)
doc.add_paragraph(
    'Orbit runs real Chrome extensions - an ad blocker, a password manager, a note clipper - and '
    'they apply to everyone, because there is one browser.'
)
para = doc.add_paragraph()
para.add_run('Using one: ').bold = True
para.add_run(
    'open ⋮ menu, Extensions. Each installed extension is listed, and each has a button that opens '
    'its own page. If it also has a settings page, an Options button appears next to it.'
)
callout('Why it opens as a tab',
        'In a normal browser an extension\'s popup is a small floating window. Chromium draws '
        'those as native desktop windows, outside the page area that is streamed to you - so a '
        'floating popup could never appear on your screen. Orbit opens the extension\'s page as a '
        'tab instead. Same page, same buttons, just docked in a tab.', fill='E7F3EB')
doc.add_paragraph(
    'One may already be there: Orbit offers a default extension when it starts, and whoever set it '
    'up decided whether to accept it. If the Extensions panel is empty, nobody installed anything '
    'yet - ask an admin.'
)
para = doc.add_paragraph()
para.add_run('Installing one (admins only): ').bold = True
para.add_run(
    'in the same panel, paste a Chrome Web Store link - or just the extension id - into "Store URL '
    'or extension id" and click Add. Then restart the browser, because Chromium only reads its '
    'extensions when it starts.'
)
doc.add_paragraph(
    'If you are not an admin you will not see the Add box; ask an admin for anything the team needs.'
)
callout('Expect some not to work',
        'These are real Chrome extensions in a real Chromium, and most behave normally - their '
        'keyboard shortcuts work too. But some will not. An extension whose popup is drawn as a '
        'native panel with no web page behind it has nothing for Orbit to open, and the panel says '
        '"no page". Anything that expects to talk to a program installed on your own computer, or '
        'that wants you to sign in to the browser itself, will not behave as it does at home. '
        'Trying one costs nothing - install it, look, and remove it if it does not work.')

doc.add_heading('2.14  Who else is here, and how it is doing', level=2)
bullets([
    '**The bottom strip** shows who is signed in right now, and whether your connection is '
    'healthy, idle or reconnecting. If it says reconnecting, it is already retrying - you do not '
    'have to reload.',
    '**Other people\'s cursors** appear over the page when you are looking at the same tab, so '
    'you can see what a colleague is pointing at.',
    '**Performance metrics** (⋮ menu) turns on a live read-out of frame rate and delay, if you '
    'want to see why something feels slow.',
    '**Appearance** (⋮ menu) cycles light, dark and follow-my-system, and remembers your choice.',
])

doc.add_heading('2.15  The one thing to remember', level=2)
callout('Important',
        'Logins are shared. If you sign in to Gmail, Jira or anything else inside Orbit, everyone '
        'using Orbit is signed in to that account too - and stays signed in after you close your '
        'tab. Never sign in to a personal account, or to anything you would not hand to the whole '
        'team.', fill='FBE2E2')
doc.add_paragraph(
    'That sharing is the point of the tool - one set of team logins everybody can use, without '
    'passing passwords around. It just means treating Orbit as a shared room rather than your own '
    'desk. Two habits help:'
)
bullets([
    'Sign out of a website inside Orbit when the team no longer needs it.',
    'Use your own browser for anything personal.',
])

# =============================================================================
part(3, 'Running it', 'For whoever hosts Orbit: settings, extensions, outside access, upkeep.')
# =============================================================================

doc.add_heading('3.1  Settings, in plain words', level=2)
doc.add_paragraph(
    'All settings live in the .env file in your Orbit folder - written for you on the first run. '
    'Change a line, then run ./orbit up to apply it. None of these are required; each has a '
    'sensible default, and ./orbit env shows what is in use with secrets hidden.'
)
table(['Setting', 'What it does', 'Default'], [
    ['APP_PORT', 'Which port Orbit is served on. Change it if 3030 is taken.', '3030'],
    ['HOME_URL', 'The page new tabs open on.', 'Google'],
    ['VIEWPORT_WIDTH / _HEIGHT', 'The page size everyone sees, in pixels. Bigger fits more on screen but makes text smaller.', '1920 x 1080'],
    ['DEVICE_SCALE_FACTOR', 'Sharpness. Same layout, drawn with more pixels. 1 to 3, decimals allowed. Only helps on high-resolution (Retina) screens, and uses more bandwidth.', '1'],
    ['STREAM_QUALITY', 'Picture quality, 1-100. Lower means less bandwidth.', '80'],
    ['MAX_FPS', 'How smooth motion is. Lower uses less processor.', '30'],
    ['PERSIST_SESSION_COOKIES', 'Keep everyone signed in to websites when Orbit restarts.', 'true'],
    ['TAB_OWNERSHIP', 'Tabs belong to whoever opened them. Off means anyone can type in any tab.', 'true'],
    ['MAX_TABS / MAX_USERS', 'Limits, so one busy day cannot overwhelm the machine.', '20 / 50'],
    ['CPU_LIMIT / MEMORY_LIMIT', 'How much of the host computer Orbit may use.', '4.0 / 4g'],
    ['CHROMIUM_TIMEZONE', 'What time zone websites think you are in.', 'UTC'],
], widths=[1.85, 3.6, 1.05])
callout('Leave this one alone',
        'DEVTOOLS_ENABLED gives admins developer tools inside the shared browser. Anyone who opens '
        'it can read what any page holds, including other people\'s signed-in sessions. It ships '
        'off, and off is the right default.')

doc.add_heading('3.2  Installing extensions', level=2)
doc.add_paragraph(
    'Extensions are installed once and appear for everyone. Two ways in, both admin-only: the '
    'Extensions panel (⋮ menu, Extensions) has an Add box, and the Admin panel has a fuller '
    'Extensions section for removing them.'
)
steps([
    'Paste a Chrome Web Store link, or just the extension id, into **"Store URL or extension id"**.',
    'Click **Add**.',
    'Restart the browser: **Admin panel, Browser, restart**.',
])
callout('The restart is not optional',
        'Chromium reads its extensions only when it starts, so a newly added extension will not '
        'appear until you restart the browser. Everyone stays signed in to their websites across '
        'that restart, as long as PERSIST_SESSION_COOKIES is on.')
doc.add_paragraph(
    'How people then use an extension - and why its popup opens as a tab - is section 2.13, which '
    'also explains why some extensions will not work at all. Try them: nothing breaks, and '
    'removing one is a click in the Admin panel plus a restart.'
)

para = doc.add_paragraph()
para.paragraph_format.space_before = Pt(8)
para.add_run('The default extension').bold = True
doc.add_paragraph(
    'Orbit ships with one: DOM Heist, from github.com/Astro-Dude/VibeExtract. Every ./orbit up and '
    './orbit restart fetches the latest copy and, the first time, asks whether to install it:'
)
code([
    '> Orbit ships with a default extension:',
    '    DOM Heist 3.2.0',
    '    from https://github.com/Astro-Dude/VibeExtract.git',
    '    it can: activeTab, scripting, webNavigation, storage, downloads, clipboardWrite',
    '  install it? [Y/n]',
])
bullets([
    'You are asked **once**. The answer is remembered, so starting Orbit never nags.',
    'After that, an update published upstream is installed on the next start without asking, and '
    'the browser only restarts when something actually changed.',
    'Say no and it is never installed or asked about again. To be asked afresh, delete '
    '.orbit-cache/extensions/<name>.decision',
    'To remove it later: ./orbit ext rm <id>, then restart the browser.',
])
doc.add_paragraph('To answer in advance, or to change which extensions are offered:')
table(['Setting', 'Effect'], [
    ['ORBIT_INSTALL_DEFAULT_EXTENSIONS=yes', 'Install without asking - for scripts and unattended machines.'],
    ['ORBIT_INSTALL_DEFAULT_EXTENSIONS=no', 'Never install, never ask.'],
    ['ORBIT_DEFAULT_EXTENSIONS="<git-url> ..."', 'Which extensions are offered. Empty means none at all.'],
], widths=[2.9, 3.6])
callout('It never blocks a start',
        'An unreachable repository, a missing git, or no terminal to ask at prints a warning and '
        'Orbit starts anyway. Nothing is installed without an answer - if there is nowhere to ask, '
        'it is skipped.')

doc.add_heading('3.3  Access from outside the network, with ngrok', level=2)
doc.add_paragraph(
    'By default Orbit is reachable only on your own network. ngrok gives it a temporary public web '
    'address, so somebody at home or in another office can use it.'
)
steps([
    'Install ngrok from ngrok.com (on a Mac: brew install ngrok).',
    'Sign up free, copy your authtoken from their dashboard, and run it once: '
    'ngrok config add-authtoken YOUR_TOKEN',
    'With Orbit already running, start the tunnel: **ngrok http 3030**',
    'It prints a public address such as https://random-words-1234.ngrok-free.app',
    'Tell Orbit it is behind HTTPS by adding three lines to .env (below), then run '
    '**./orbit up**',
    'Share the https:// link, and each person\'s username and password.',
])
code([
    'TRUSTED_ORIGINS=https://random-words-1234.ngrok-free.app',
    'SECURE_COOKIES=true',
    'TRUST_PROXY=true',
])
table(['Line', 'Why it is needed'], [
    ['TRUSTED_ORIGINS', 'Tells Orbit this public address is legitimately yours. Must match exactly, including https:// and no trailing slash.'],
    ['SECURE_COOKIES', 'Sends the sign-in cookie only over HTTPS.'],
    ['TRUST_PROXY', 'ngrok sits in front, so Orbit reads the real visitor address from it.'],
], widths=[1.7, 4.8])
callout('This breaks the local address',
        'Turning on SECURE_COOKIES stops the plain http://192.168... address from working: browsers '
        'refuse to store a secure cookie over plain HTTP, so nobody on your own network can sign in '
        'there any more. Everyone uses the https:// link instead. To go back to local-only, remove '
        'those three lines and run ./orbit up again.', fill='FBE2E2')
doc.add_paragraph('Before you publish it, know these:')
bullets([
    '**The free address changes** each time you restart ngrok, and TRUSTED_ORIGINS has to be '
    'updated to match. ngrok\'s free plan includes one reserved domain - set it up in their '
    'dashboard and the address stops moving.',
    '**It will feel slower.** Everything you click travels to the host computer and back. On your '
    'own network that is a few milliseconds; from another country it can be 200 ms or more - fine '
    'for reading and clicking, noticeably laggy while typing.',
    '**Anyone with the link can try passwords against it.** Use strong ones, remove people you no '
    'longer want (Admin panel, Users), and stop ngrok when you are done for the day.',
    '**Remember the shared logins.** A public link is a public door to every account the browser '
    'is signed into. Sign out of anything sensitive first.',
])

doc.add_heading('3.4  Everyday commands', level=2)
doc.add_paragraph('Run these in your orbit folder.')
table(['Command', 'Does'], [
    ['./orbit up', 'Start Orbit, and apply any .env changes.'],
    ['./orbit status', 'Is it running and healthy? Plus live figures and the addresses.'],
    ['./orbit url', 'Just the addresses, to paste to your team.'],
    ['./orbit restart', 'Restart the browser, keeping everyone signed in.'],
    ['./orbit down', 'Stop it, keep everything.'],
    ['./orbit logs 50', 'Recent messages - useful when something is wrong.'],
    ['git pull && ./orbit up', 'Update to the newest version of Orbit.'],
    ['./orbit down --wipe', 'WARNING: deletes everything - users, logins, history.'],
], widths=[2.6, 3.9])

doc.add_heading('3.5  Where the data lives, and backups', level=2)
doc.add_paragraph(
    'Logins, bookmarks, history and accounts live in a Docker volume named orbit-data, separate '
    'from the browser itself - so updating Orbit keeps all of it. To save and restore a copy:'
)
code(['./orbit backup', './orbit restore orbit-backup-20260831-120000.tar.gz'])
doc.add_paragraph(
    'Backup writes orbit-backup-<date>.tar.gz into the current folder. It stops the browser for a '
    'few seconds first so Chromium finishes writing its profile, then starts it again - people '
    'using Orbit will see a brief interruption. Keep the file somewhere sensible: it contains the '
    'browser\'s saved logins.'
)

# =============================================================================
part(4, 'When something goes wrong', 'Most problems are in this table. The three after it are the ones that need more than a line.')
# =============================================================================

doc.add_heading('4.1  Quick table', level=2)
table(['What you see', 'What to do'], [
    ['docker: command not found', 'Docker is not installed, or Docker Desktop is not open yet.'],
    ['"set SESSION_SECRET in .env" when starting', 'Your .env is missing, empty, or not in the folder you ran the command from.'],
    ['"port is already allocated"', 'Something else uses 3030. Put APP_PORT=3040 in .env, start again, and use :3040 in the address.'],
    ['Status stuck on "starting"', 'Give it 45 seconds. Then: ./orbit logs 50'],
    ['The container keeps restarting', 'Usually not enough memory. Raise MEMORY_LIMIT, or close other heavy apps on the host.'],
    ['Works on the host, not from other devices', 'See section 4.3.'],
    ['Chrome on a Mac: ERR_ADDRESS_UNREACHABLE', 'See section 4.2.'],
    ['The page area is blank, black or frozen', 'Click Reload in Orbit\'s toolbar. If it stays blank, restart the browser: Admin panel, Browser.'],
    ['Text looks slightly soft on a Retina screen', 'Set DEVICE_SCALE_FACTOR=1.5 in .env, then ./orbit up.'],
    ['Slow or jerky for everyone', 'Lower STREAM_QUALITY to 70 and MAX_FPS to 30. Fewer open tabs helps most of all.'],
    ['Everyone signed out of websites after a restart', 'Set PERSIST_SESSION_COOKIES=true in .env.'],
    ['Installed an extension but cannot see it', 'Restart the browser: Admin panel, Browser. Extensions only load at startup.'],
    ['My own browser\'s shortcut fires instead of Orbit\'s', 'Turn on Capture keyboard (Option / Alt + K).'],
    ['I asked for control and nothing happened', 'The owner may be away. Open your own tab to the same page, or ask them directly.'],
    ['I cannot type in a tab', 'Somebody else owns it - use "Ask <name> for control". Viewers can never type; ask an admin to change your role.'],
    ['Cannot sign in over the ngrok link', 'See section 4.4.'],
    ['Forgot the admin password', 'Sign in as another admin and reset it in Admin panel, Users. With no admin left, ./orbit down --wipe starts fresh - and deletes everything.'],
], widths=[2.5, 4.0])

doc.add_heading('4.2  Chrome on a Mac says "This site can\'t be reached"', level=2)
doc.add_paragraph(
    'You open the Orbit address and Chrome reports ERR_ADDRESS_UNREACHABLE - but the same link '
    'works in Safari.'
)
doc.add_paragraph(
    'Nothing is wrong with Orbit or your network. Recent versions of macOS ask each app for '
    'permission to talk to devices on the local network, and Chrome has been refused. macOS blocks '
    'it before it reaches the network, so Chrome reports an unreachable address rather than '
    '"permission denied" - which is why it looks like a network fault. Safari does not need that '
    'permission, so it keeps working.'
)
doc.add_paragraph('Try these in order:')
steps([
    'Open **System Settings**, go to **Privacy & Security**, then **Local Network**, and switch '
    '**Google Chrome** on. Then quit Chrome completely with **Command+Q** - closing the window is '
    'not enough - and open it again.',
    'If Chrome is not in that list, or it is already on and still fails: in Chrome go to '
    '**chrome://settings/content/localNetworkAccess** and choose "Don\'t allow sites to connect to '
    'any device on your local network". Wait about ten seconds, then switch to "Sites can ask..." '
    'and click **Allow** when the prompt appears. This makes macOS ask again.',
    'If both fail, the permission is stuck in a macOS cache. In Terminal run '
    '**tccutil reset All com.google.Chrome**, then **restart the Mac**. The restart is required - '
    'without it the reset has no effect.',
])
doc.add_paragraph(
    'While you are there: check that no VPN or proxy extension is switched on in Chrome only. That '
    'produces the same error for a different reason.'
)

doc.add_heading('4.3  Other devices cannot reach it', level=2)
steps([
    '**Same network?** A phone on mobile data, or a guest Wi-Fi that isolates devices, will not '
    'reach it.',
    '**Right address?** Check it again with the command in section 1.4 - it changes when the host '
    'computer reconnects to the network.',
    '**Firewall on the host?** Allow incoming connections on port 3030. macOS: System Settings, '
    'Network, Firewall, Options, allow Docker. Windows: Windows Defender Firewall, Allow an app, '
    'Docker Desktop. Linux: sudo ufw allow 3030/tcp',
    '**Still nothing?** On the host, curl -s localhost:3030/api/health should print '
    '{"status":"running",...}. If it does, Orbit is fine and the problem is between the devices.',
])

doc.add_heading('4.4  ngrok: the link opens but I cannot sign in', level=2)
doc.add_paragraph('Two usual causes:')
bullets([
    '**TRUSTED_ORIGINS does not match.** It must be the exact address, including https:// and with '
    'no trailing slash. If your ngrok address changed, update the line and run '
    './orbit up.',
    '**You are on the plain http:// address while SECURE_COOKIES is true.** The browser will not '
    'keep the sign-in cookie. Use the https:// link, or remove those three lines to go back to '
    'local use.',
])
doc.add_paragraph(
    'If the sign-in page will not load at all, check ngrok is still running - the free tunnel '
    'closes when its window is closed or the host laptop sleeps.'
)

doc.add_heading('4.5  Who to ask', level=2)
doc.add_paragraph('Some things only an admin can do. Ask yours for:')
bullets([
    'A new account, a password reset, or a change to what you are allowed to do.',
    'Installing an extension for the team.',
    'Restarting the shared browser.',
    'Changing picture quality, sharpness, or the page size everyone sees.',
])
last = doc.add_paragraph()
r = last.add_run(
    'Not sure which role you have? Try opening a tab - viewers cannot.'
)
r.italic = True
r.font.size = Pt(9.5)

page_numbers()
doc.save(OUT)
print('written:', OUT)
