"""Build the end-user usage guide as a Word document.

    pip install python-docx
    python3 scripts/make-usage-docx.py docs/Orbit-Usage-Guide.docx

Kept so the .docx can be regenerated when the UI changes, rather than being a
binary nobody can edit. The prose here and in docs/team-guide.md cover the same
ground for different readers - change both, or neither.

Content mirrors docs/team-guide.md sections 3-4 and the user-facing half of the
troubleshooting, but aimed at someone who is only *using* Orbit: no Docker, no
.env, no host setup. Every label quoted here was read out of the UI source.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

OUT = sys.argv[1]
doc = Document()

# --- document setup ---------------------------------------------------------
p = doc.core_properties
p.title = 'Orbit - How to Use It'
p.subject = 'Using the shared team browser'
p.author = 'the-ashutosh'
p.comments = 'End-user guide for Orbit, the shared multi-user browser.'

for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8)
    s.left_margin = s.right_margin = Inches(0.9)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def callout(label, text, fill='FFF3CD'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(10)
    para.paragraph_format.left_indent = Inches(0.12)
    r = para.add_run(f'{label}  ')
    r.bold = True
    para.add_run(text)
    shade(para, fill)
    return para


def bullets(items, style='List Bullet'):
    for it in items:
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(3)
        # **bold** segments
        for i, chunk in enumerate(it.split('**')):
            run = para.add_run(chunk)
            run.bold = i % 2 == 1


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    for style in ('Light List Accent 1', 'Light Grid Accent 1', 'Table Grid'):
        try:
            t.style = style
            break
        except KeyError:
            continue
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            for j, chunk in enumerate(str(val).split('**')):
                run = para.add_run(chunk)
                run.bold = j % 2 == 1
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def page_numbers():
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run('Orbit - How to Use It    ').font.size = Pt(8)
    run = para.add_run()
    run.font.size = Pt(8)
    for kind, text in (('begin', None), (None, 'PAGE'), ('end', None)):
        if text:
            el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve'); el.text = text
        else:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), kind)
        run._r.append(el)


# --- title ------------------------------------------------------------------
title = doc.add_heading('Orbit', level=0)
sub = doc.add_paragraph()
r = sub.add_run('How to use the shared team browser')
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.paragraph_format.space_after = Pt(14)

doc.add_paragraph(
    'Orbit is one real browser running on one computer in your team, shared with everyone. '
    'You open it in your own browser and work in your own tabs - several people browse at the '
    'same time, independently, in the same shared browser.'
)
doc.add_paragraph(
    'It is not screen sharing. Nobody has to present, and you are not watching someone else '
    'move their mouse. The one thing that IS shared is logins: sign in to a website once and '
    'everyone using Orbit is signed in to it.'
)

# --- 1. getting in ----------------------------------------------------------
doc.add_heading('1. Getting in', level=1)
doc.add_paragraph('Ask whoever set Orbit up for three things:')
bullets([
    '**The address** - something like http://192.168.1.42:3030, or an https:// link if you are outside the office.',
    '**Your username**.',
    '**Your password**.',
])
doc.add_paragraph(
    'Open the address in any browser, on any device - laptop, phone or tablet. Type your username '
    'and password and click Login. Nothing to install.'
)
callout('Note',
        'You must be on the same Wi-Fi or network as the computer running Orbit, unless you were '
        'given an https:// link. Mobile data will not reach a 192.168.x.x address.')

# --- 2. the window ----------------------------------------------------------
doc.add_heading('2. The window', level=1)
table(['Where', 'What it is'], [
    ['Top strip', 'The tabs. Click one to switch, click + for a new one, x to close.'],
    ['Toolbar below it', 'Back, Forward, Reload, the address bar, and a star to bookmark the page.'],
    ['⋮ button, top right', 'The menu: zoom, full screen, bookmarks, history, downloads, extensions, appearance, sign out.'],
    ['Bottom strip', 'Who else is online and how the connection is doing.'],
], widths=[1.6, 4.9])
doc.add_paragraph(
    'The address bar works the way you expect: type a web address, or type words to search. '
    'It also suggests pages that have been visited before.'
)

# --- 3. tabs and control ----------------------------------------------------
doc.add_heading('3. Tabs belong to whoever opens them', level=1)
doc.add_paragraph(
    'This is the one rule worth understanding. When you open a tab, that tab is yours. Other '
    'people can see it, but they cannot type or click in it - so nobody fills in your form or '
    'clicks Send on your email by accident.'
)
doc.add_paragraph('To use a tab somebody else opened:')
bullets([
    'Open the tab. The toolbar shows a button: **"Ask <name> for control"**.',
    'Click it. The button changes to **"Asked <name>..."** while you wait.',
    'The owner sees a small box in the corner of that same tab: **"Can I drive this tab?"** '
    'with two buttons, **Give control** and **Keep it to myself**.',
    'If they give control, you can type and click in that tab too.',
])
doc.add_paragraph(
    'Nothing is lost if they say no, or are away from their desk - you can still read the page, '
    'and you can always open your own tab and go to the same site.'
)
callout('Tip',
        'Working on something you do not want touched? Just keep it in a tab you opened. That is '
        'already private to you, with no extra steps.', fill='E7F3EB')

# --- 4. shortcuts -----------------------------------------------------------
doc.add_heading('4. Keyboard shortcuts', level=1)
doc.add_paragraph(
    'Orbit\'s own shortcuts use Option on a Mac and Alt on Windows and Linux. They deliberately '
    'avoid Ctrl and Command, so they never clash with your own browser\'s shortcuts.'
)
table(['Press', 'Does'], [
    ['Option / Alt + T', 'New tab'],
    ['Option / Alt + W', 'Close the tab'],
    ['Option / Alt + D', 'Jump to the address bar'],
    ['Option / Alt + F', 'Full screen'],
    ['Option / Alt + K', 'Capture keyboard (see below)'],
    ['Option / Alt + left / right arrow', 'Back / Forward'],
    ['Option / Alt + 1 ... 9', 'Switch to that tab'],
], widths=[2.3, 4.2])
doc.add_paragraph()
h = doc.add_paragraph(); h.add_run('Copy and paste are normal.').bold = True
doc.add_paragraph(
    'Command+C / Command+V on a Mac, Ctrl+C / Ctrl+V on Windows. They work in both directions: '
    'copy in Orbit and paste into your own documents, or copy from your computer and paste into '
    'a page in Orbit.'
)

# --- 5. capture keyboard ----------------------------------------------------
doc.add_heading('5. Capture keyboard', level=1)
doc.add_paragraph(
    'Some keys belong to your own browser. Press Ctrl+T (or Command+T) and YOUR browser opens a '
    'tab - Orbit never sees it. Turn on Capture keyboard and those keys go to Orbit instead.'
)
bullets([
    'Turn it on from the ⋮ menu, or with **Option / Alt + K**.',
    'It applies to **that one tab only**, and only while you are looking at it.',
    'A small badge appears so you know it is on. Turn it off the same way.',
])
doc.add_paragraph(
    'Use it when you want a page to feel like a normal browser tab. Leave it off the rest of the '
    'time, so your own browser shortcuts keep working.'
)

# --- 6. the menu ------------------------------------------------------------
doc.add_heading('6. The rest of the menu', level=1)
table(['Menu item', 'What it does'], [
    ['Zoom', 'Makes the page bigger or smaller. At the top of the menu.'],
    ['New tab / Duplicate tab', 'As it sounds.'],
    ['Full screen', 'The page fills your screen. Orbit\'s tabs and toolbar stay visible.'],
    ['Bookmarks', 'Saved pages - shared with the whole team.'],
    ['History', 'Pages visited in the shared browser.'],
    ['Downloads', 'Files downloaded inside Orbit, ready for you to save to your own computer.'],
    ['Extensions', 'Chrome extensions installed for everyone. You can open an extension\'s popup here. Installing is an admin job.'],
    ['Appearance', 'Light, dark, or follow your system.'],
    ['Performance metrics', 'Live frame rate and delay, if you are curious.'],
    ['Sign out', 'Ends your session only. Everyone else keeps working.'],
], widths=[1.9, 4.6])

# --- 7. shared logins -------------------------------------------------------
doc.add_heading('7. The one thing to remember', level=1)
callout('Important',
        'Logins are shared. If you sign in to Gmail, Jira or anything else inside Orbit, everyone '
        'using Orbit is signed in to that account too - and stays signed in after you close your '
        'tab. Never sign in to a personal account, and never to anything you would not hand to '
        'the whole team.', fill='FBE2E2')
doc.add_paragraph(
    'That sharing is the point of the tool - one set of team logins everybody can use, without '
    'passing passwords around. It just means treating Orbit as a shared room, not your own desk.'
)
doc.add_paragraph('Two habits that help:')
bullets([
    'Sign out of a website inside Orbit when the team no longer needs it.',
    'Use your own browser for anything personal. Orbit is for shared work.',
])

# --- 8. troubleshooting -----------------------------------------------------
doc.add_heading('8. If something goes wrong', level=1)
table(['What you see', 'What to do'], [
    ['The address will not open at all', 'Check you are on the same Wi-Fi as the computer running Orbit. Phones on mobile data cannot reach it. Then re-check the address with whoever set it up - it changes if that computer reconnects to the network.'],
    ['Chrome on a Mac says "ERR_ADDRESS_UNREACHABLE" but Safari works', 'A macOS permission, not a network fault. See the next section.'],
    ['The page area is blank, black, or frozen', 'Click Reload in Orbit\'s toolbar. If it stays blank, ask an admin to restart the browser.'],
    ['Everything feels slow or jerky', 'Close tabs you are not using - every open, animating tab costs the shared machine. If it is slow for everyone, tell an admin: there are quality settings they can lower.'],
    ['Text looks slightly soft', 'Expected on a high-resolution screen at default settings. An admin can raise the sharpness setting.'],
    ['I was signed out of a website', 'The shared browser was restarted. Sign in again - the team keeps that login.'],
    ['My keyboard shortcut did nothing', 'Orbit uses Option / Alt, not Ctrl or Command. For a page that needs Ctrl shortcuts, turn on Capture keyboard (Option / Alt + K).'],
    ['I asked for control and nothing happened', 'The owner may be away. Open your own tab to the same page instead, or ask them directly.'],
    ['I cannot type in a tab', 'Somebody else owns it - use "Ask <name> for control". If you are a viewer, you can watch but never type; ask an admin to change your role.'],
], widths=[2.2, 4.3])

doc.add_heading('Chrome on a Mac: "This site can\'t be reached"', level=2)
doc.add_paragraph(
    'You open the Orbit address and Chrome says the site is unreachable, with the code '
    'ERR_ADDRESS_UNREACHABLE - but the very same link works in Safari.'
)
doc.add_paragraph(
    'Nothing is wrong with Orbit or your network. Recent versions of macOS ask each app for '
    'permission to talk to devices on your local network, and Chrome has been refused. macOS '
    'blocks it before it reaches the network, so Chrome reports it as an unreachable address '
    'instead of "permission denied" - which is why it looks like a network problem. Safari does '
    'not need that permission, so it keeps working.'
)
doc.add_paragraph('Try these in order:')
for i, step in enumerate([
    'Open System Settings, go to Privacy & Security, then Local Network, and switch Google Chrome on. '
    'Then quit Chrome completely with Command+Q - closing the window is not enough - and open it again.',
    'If Chrome is not in that list, or it is already switched on and still fails: in Chrome, go to '
    'chrome://settings/content/localNetworkAccess and choose "Don\'t allow sites to connect to any '
    'device on your local network". Wait about ten seconds, then switch it to "Sites can ask..." and '
    'click Allow when the prompt appears. This makes macOS ask again.',
    'If both fail, the permission is stuck in a macOS cache. Open Terminal, run '
    'tccutil reset All com.google.Chrome and then restart your Mac. The restart is required - '
    'without it the reset has no effect.',
], start=1):
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.space_after = Pt(5)
    para.add_run(step)
doc.add_paragraph(
    'While you are checking: make sure no VPN or proxy extension is switched on in Chrome only. '
    'That produces the same error for a different reason.'
)

# --- 9. asking for help -----------------------------------------------------
doc.add_heading('9. Who to ask', level=1)
doc.add_paragraph('Some things only an admin can do. Ask yours for:')
bullets([
    'A new account, a password reset, or a change to what you are allowed to do.',
    'Installing a Chrome extension for the team.',
    'Restarting the shared browser.',
    'Changing picture quality, sharpness or the page size everyone sees.',
])
doc.add_paragraph()
last = doc.add_paragraph()
r = last.add_run('Roles: an admin manages everything, a user browses and opens tabs normally, and '
                 'a viewer can watch but never type. If you are not sure which you are, try opening '
                 'a tab - viewers cannot.')
r.font.size = Pt(9.5)
r.italic = True

page_numbers()
doc.save(OUT)
print('written:', OUT)
